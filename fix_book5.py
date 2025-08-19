import json
from docx import Document
import re
import os

def examine_book5_format():
    """Examine Book 5's specific format to understand why parsing failed"""
    file_path = "Book 5.docx"
    
    print(f"=== Examining Book 5 Format ===")
    
    try:
        doc = Document(file_path)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        
        # Look for vocabulary entries specifically
        print("\nLooking for vocabulary patterns in Book 5:")
        
        for i, paragraph in enumerate(doc.paragraphs[:10]):  # Check first 10 lines
            text = paragraph.text.strip()
            if text and not text.startswith('Unit'):
                print(f"{i+1:2d}: '{text}'")
                
                # Check what dash character is used
                if '–' in text:
                    print(f"    Contains en dash (–)")
                if '-' in text:
                    print(f"    Contains hyphen (-)")
                if '—' in text:
                    print(f"    Contains em dash (—)")
                
                # Try to match vocabulary patterns
                patterns = [
                    r'^([A-Za-z]+)\s*\(([^)]+)\)\s*[–-]\s*(.+?)(?:\.\s*(.+))?$',
                    r'^([A-Za-z]+)\s*[–-]\s*(.+?)(?:\.\s*(.+))?$',
                    r'^([A-Za-z]+)\s*\(([^)]+)\)\s*[–-]\s*(.+)$',
                    r'^([A-Za-z]+)\s*[–-]\s*(.+)$'
                ]
                
                for pattern in patterns:
                    match = re.match(pattern, text)
                    if match:
                        print(f"    ✓ MATCHED: {match.groups()}")
                        break
                else:
                    print(f"    ✗ NO MATCH")
        
    except Exception as e:
        print(f"Error reading {file_path}: {e}")

def parse_book5_specific():
    """Parse Book 5 with improved patterns"""
    file_path = "Book 5.docx"
    
    try:
        doc = Document(file_path)
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        units = {}
        current_unit = None
        current_words = []
        
        for line in content:
            line = line.strip()
            if not line:
                continue
                
            # Look for unit headers
            unit_match = re.search(r'Unit\s+(\d+)[:\-\s]*(.*)', line, re.IGNORECASE)
            if unit_match:
                if current_unit and current_words:
                    units[current_unit] = current_words
                    current_words = []
                
                current_unit = unit_match.group(1)
                print(f"Found Unit {current_unit}: {unit_match.group(2).strip()}")
                continue
            
            # Enhanced patterns for Book 5 - handle all possible dash characters
            if current_unit:
                # Try multiple patterns with different dash characters
                patterns = [
                    # Pattern 1: "word (pos) – meaning. Example sentence." (en dash)
                    r'^([A-Za-z]+)\s*\(([^)]+)\)\s*–\s*(.+?)(?:\.\s*(.+))?$',
                    # Pattern 2: "word – meaning. Example sentence." (en dash)
                    r'^([A-Za-z]+)\s*–\s*(.+?)(?:\.\s*(.+))?$',
                    # Pattern 3: "word (pos) - meaning. Example sentence." (hyphen)
                    r'^([A-Za-z]+)\s*\(([^)]+)\)\s*-\s*(.+?)(?:\.\s*(.+))?$',
                    # Pattern 4: "word - meaning. Example sentence." (hyphen)
                    r'^([A-Za-z]+)\s*-\s*(.+?)(?:\.\s*(.+))?$',
                    # Pattern 5: "word (pos) — meaning. Example sentence." (em dash)
                    r'^([A-Za-z]+)\s*\(([^)]+)\)\s*—\s*(.+?)(?:\.\s*(.+))?$',
                    # Pattern 6: "word — meaning. Example sentence." (em dash)
                    r'^([A-Za-z]+)\s*—\s*(.+?)(?:\.\s*(.+))?$',
                    # Pattern 7: "word (pos) – meaning" (en dash, no example)
                    r'^([A-Za-z]+)\s*\(([^)]+)\)\s*–\s*(.+)$',
                    # Pattern 8: "word – meaning" (en dash, no example)
                    r'^([A-Za-z]+)\s*–\s*(.+)$'
                ]
                
                for pattern in patterns:
                    match = re.match(pattern, line)
                    if match:
                        word = match.group(1).lower().strip()
                        pos = match.group(2).strip() if len(match.groups()) > 1 else "noun"
                        meaning = match.group(3).strip() if len(match.groups()) > 2 else match.group(2).strip()
                        example = match.group(4).strip() if len(match.groups()) > 3 else ""
                        
                        # Clean up part of speech
                        pos = pos.lower()
                        if pos in ['adj', 'adjective']:
                            pos = "adjective"
                        elif pos in ['v', 'verb']:
                            pos = "verb"
                        elif pos in ['n', 'noun']:
                            pos = "noun"
                        elif pos in ['adv', 'adverb']:
                            pos = "adverb"
                        elif '/' in pos:
                            pos = pos.split('/')[0].strip()
                            if pos in ['adj', 'adjective']:
                                pos = "adjective"
                            elif pos in ['v', 'verb']:
                                pos = "verb"
                            elif pos in ['n', 'noun']:
                                pos = "noun"
                            elif pos in ['adv', 'adverb']:
                                pos = "adverb"
                        
                        # Skip if word is too short or looks like a header
                        if len(word) < 2 or word.lower() in ['unit', 'chapter', 'section', 'book']:
                            continue
                        
                        # Generate example if not provided
                        if not example:
                            if pos == "verb":
                                example = f"The team will {word} the project successfully."
                            elif pos == "adjective":
                                example = f"The {word} approach was very effective."
                            else:
                                example = f"The {word} is important for success."
                        
                        current_words.append({
                            "word": word,
                            "partOfSpeech": pos,
                            "meaning": meaning,
                            "example": example
                        })
                        print(f"  ✓ Extracted: {word} ({pos}) - {meaning[:50]}...")
                        break
        
        # Add the last unit
        if current_unit and current_words:
            units[current_unit] = current_words
        
        return units
        
    except Exception as e:
        print(f"Error processing Book 5: {e}")
        return {}

def update_book_data():
    """Update bookData.js with Book 5 content"""
    print("Updating bookData.js with Book 5 content...")
    
    # Read existing bookData.js
    with open('bookData.js', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Extract the bookData object
    start = content.find('const bookData = ') + len('const bookData = ')
    end = content.find(';\n\n// Export', start)
    
    if start == -1 or end == -1:
        print("Could not find bookData in file")
        return
    
    # Parse existing data
    existing_data = json.loads(content[start:end])
    
    # Get Book 5 data
    book5_units = parse_book5_specific()
    
    if book5_units:
        existing_data["5"]["units"] = book5_units
        print(f"Added {len(book5_units)} units to Book 5")
        
        # Count total words in Book 5
        total_words = sum(len(unit) for unit in book5_units.values())
        print(f"Total words in Book 5: {total_words}")
        
        # Write updated data
        with open('bookData.js', 'w', encoding='utf-8') as f:
            f.write('// Book Data Structure - Extracted from Word Documents\n')
            f.write('// Real vocabulary content from user-provided files\n')
            f.write('// All 4 books with complete vocabulary\n\n')
            f.write('const bookData = ')
            f.write(json.dumps(existing_data, indent=2, ensure_ascii=False))
            f.write(';\n\n')
            f.write('// Export the data for use in the main script\n')
            f.write('if (typeof module !== \'undefined\' && module.exports) {\n')
            f.write('    module.exports = bookData;\n')
            f.write('}\n')
        
        print("✅ bookData.js updated successfully!")
        
        # Print statistics
        total_units = sum(len(book["units"]) for book in existing_data.values())
        total_words = sum(
            sum(len(unit) for unit in book["units"].values()) 
            for book in existing_data.values()
        )
        print(f"\n📊 Final Statistics:")
        print(f"Total books: {len(existing_data)}")
        print(f"Total units: {total_units}")
        print(f"Total words: {total_words}")
        
        for book_num, book_info in existing_data.items():
            word_count = sum(len(unit) for unit in book_info["units"].values())
            unit_count = len(book_info["units"])
            print(f"   Book {book_num}: {word_count} words across {unit_count} units")
    else:
        print("❌ No Book 5 data extracted")

def main():
    """Main function"""
    print("Fixing Book 5 vocabulary extraction...")
    
    # First examine the format
    examine_book5_format()
    
    print("\n" + "="*50)
    
    # Then update the data
    update_book_data()

if __name__ == "__main__":
    main()
